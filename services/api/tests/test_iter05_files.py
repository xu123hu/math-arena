"""迭代05 文件解析去占位符测试（阶段 1.4，审计 A-P0-4）

覆盖（SSOT §5.3 决策表 / §4.2 wf_doc_understand）：
1. spark_vl 显式指定 → wf_doc_understand 云轨调用 + 输出拼装
2. spark_vl 失败 → RapidOCR 兜底链
3. mineru 未部署 → 降级 PyMuPDF
4. rapidocr 依赖缺失 → None（绝不返回占位符文本）
5. pandoc 缺失 → python-docx 兜底真实抽取（本机无 pandoc，真实路径）
6. 占位符文本彻底清除（全文件无 "[OCR 解析待接入]"/"[Pandoc 解析待接入]"）
"""

import io
import json
import pathlib
import uuid as _uuid
from unittest.mock import AsyncMock, patch

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.pool import NullPool

from app.config import settings as _settings
from app.domains.files import router as fr
from app.main import app as _app
from app.models.database import get_db as _get_db
from app.models.file import File


class _FakeStorage:
    def __init__(self, data: bytes):
        self._data = data

    def get_bytes(self, uri):
        return self._data


def _file_obj(file_type: str = "image", filename: str = "t.jpg") -> File:
    return File(file_type=file_type, filename=filename, storage_uri="fake://t")


@pytest.mark.asyncio
async def test_spark_vl_cloud_track():
    """显式 spark_vl → wf_doc_understand 调用，question_text + latex_fragments 拼装

    迭代19 契约变更：_parse_image_vision 返回 (text, confidence) 二元组，
    不再依赖已下线的星辰文件上传服务（data URI 直传）。
    """
    fake_result = {
        "question_text": "已知函数 f(x)=sin x，求周期。",
        "latex_fragments": ["f(x)=\\sin x"],
        "has_figure": False,
        "question_type": "solution",
        "confidence": 0.92,
    }
    with patch.object(fr, "get_storage", return_value=_FakeStorage(b"img-bytes")), \
         patch("app.providers.xingchen.run_workflow", new=AsyncMock(return_value=fake_result)), \
         patch.object(fr.settings, "xingchen_enabled", True):
        content, confidence = await fr._parse_image_vision(_file_obj())
    assert content is not None
    assert confidence == 0.92
    assert "已知函数 f(x)=sin x" in content
    assert "$f(x)=\\sin x$" in content


@pytest.mark.asyncio
async def test_spark_vl_disabled_returns_none():
    """星辰关闭 → 云轨直接 (None, 0.0)（调用方走本地轨）"""
    with patch.object(fr.settings, "xingchen_enabled", False):
        content, confidence = await fr._parse_image_vision(_file_obj())
    assert content is None
    assert confidence == 0.0


@pytest.mark.asyncio
async def test_spark_vl_engine_fallback_rapidocr():
    """engine=spark_vl 云轨失败 → 降级 RapidOCR（本机依赖缺失时返回 None，链路不崩）"""
    with patch.object(fr, "_parse_image_vision", new=AsyncMock(return_value=(None, 0.0))), \
         patch.object(fr, "_parse_image_rapidocr", new=AsyncMock(return_value="OCR 兜底文本")) as m:
        content = await fr._dispatch_parse(_file_obj(), "spark_vl", "question_photo")
    assert content == "OCR 兜底文本"
    m.assert_awaited_once()


@pytest.mark.asyncio
async def test_mineru_fallback_pymupdf():
    """mineru 未部署 → 降级 PyMuPDF（契约保留口径）"""
    with patch.object(fr, "_parse_pdf_pymupdf", new=AsyncMock(return_value="PDF 文本")) as m:
        content = await fr._dispatch_parse(_file_obj(file_type="pdf", filename="t.pdf"), "mineru", "kb_ingest")
    assert content == "PDF 文本"
    m.assert_awaited_once()


@pytest.mark.asyncio
async def test_rapidocr_missing_returns_none_not_placeholder():
    """rapidocr 依赖缺失 → None（绝不允许占位符文本流入资产）"""
    try:
        import rapidocr_onnxruntime  # noqa: F401
        pytest.skip("rapidocr 已安装，跳过缺失降级用例")
    except ImportError:
        pass
    with patch.object(fr, "get_storage", return_value=_FakeStorage(b"img")):
        content = await fr._parse_image_rapidocr(_file_obj())
    assert content is None


@pytest.mark.asyncio
async def test_docx_pandoc_fallback_python_docx():
    """pandoc 缺失（本机真实无 pandoc）→ python-docx 兜底真实抽取段落文本"""
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("第一段落：三角函数周期性。")
    doc.add_paragraph("第二段落：sin(x+2π)=sin x。")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    with patch.object(fr, "get_storage", return_value=_FakeStorage(docx_bytes)):
        content = await fr._parse_office_pandoc(_file_obj(file_type="docx", filename="t.docx"))
    assert content is not None
    assert "第一段落：三角函数周期性。" in content
    assert "第二段落：sin(x+2π)=sin x。" in content


@pytest.mark.asyncio
async def test_dispatch_rapidocr_engine_no_placeholder():
    """engine=rapidocr 全路径不产生占位符（依赖缺失 → None → 上层走云轨兜底/failed）"""
    with patch.object(fr, "get_storage", return_value=_FakeStorage(b"img")):
        content = await fr._dispatch_parse(_file_obj(), "rapidocr", "question_photo")
    assert content is None or "[OCR" not in (content or "")


def test_placeholder_texts_removed_from_codebase():
    """占位符文本已从解析代码中彻底清除"""
    src = pathlib.Path(fr.__file__).read_text(encoding="utf-8")
    assert "[OCR 解析待接入]" not in src
    assert "[Pandoc 解析待接入]" not in src


def test_mimo_multimodal_result_preserves_uncertainty_for_confirmation():
    """拍题视觉识别不得把不确定的题图条件静默当成确定事实。

    此测试要防止的回归：把 MiMo 的 JSON 直接降格为纯题干，丢失
    ``uncertainties`` 后让后续课堂基于臆测条件生成。
    """
    raw = """```json
    {
      "question_text": "在四棱锥 S-ABCD 中，求点 E 到平面 SBD 的距离。",
      "conditions": ["E 是 AD 的中点", "SA 垂直于平面 ABCD"],
      "diagram_entities": {"points": ["S", "A", "B", "C", "D", "E"]},
      "uncertainties": ["图中 B、C 的连线是虚线还是实线无法确认"],
      "confidence": 0.93
    }
    ```"""

    result = fr._parse_mimo_math_photo_response(raw)

    assert result["text"] == "在四棱锥 S-ABCD 中，求点 E 到平面 SBD 的距离。"
    assert result["conditions"] == ["E 是 AD 的中点", "SA 垂直于平面 ABCD"]
    assert result["confidence"] == 0.93
    assert result["needs_confirmation"] is True
    assert result["uncertainties"] == ["图中 B、C 的连线是虚线还是实线无法确认"]


def test_mimo_multimodal_result_normalizes_entity_list_for_downstream_consumers():
    """视觉模型偶尔以列表返回图形实体，后续课堂链路仍应得到稳定对象结构。

    此测试要防止的回归：真实拍题响应中的 ``diagram_entities`` 类型漂移，
    使教材关联或 GeoGebra 读取元数据时崩溃。
    """
    raw = json.dumps(
        {
            "question_text": "如图，求二面角的余弦值。",
            "conditions": ["AE 垂直 CD"],
            "diagram_entities": ["点 A、B、C、D、E", "平面 AEC 垂直平面 ABC"],
            "uncertainties": [],
            "confidence": 0.95,
        }
    )

    result = fr._parse_mimo_math_photo_response(raw)

    assert result["diagram_entities"] == {
        "items": ["点 A、B、C、D、E", "平面 AEC 垂直平面 ABC"]
    }


def test_mimo_math_photo_request_uses_v25_image_input_contract():
    """拍题必须请求配置的视觉模型与图片输入契约，而不是纯文本模型。

    此测试要防止的回归：视觉请求退回不支持图片的文本模型，导致 API 返回
    "No endpoints found that support image input"，课堂又生成空白页面。
    视觉模型由配置驱动（历史上是 MiMo mimo-v2.5，现为智谱 glm-4v-flash）。
    """
    from app.config import settings as app_settings

    payload = fr._build_mimo_math_photo_payload("data:image/png;base64,ZmFrZQ==")

    assert payload["model"] == app_settings.mimo_vision_model
    assert payload["max_tokens"] <= 1024  # glm-4v-flash 输出上限，超限会被 1210 拒绝
    content = payload["messages"][0]["content"]
    assert content[0]["type"] == "text"
    assert content[1] == {
        "type": "image_url",
        "image_url": {"url": "data:image/png;base64,ZmFrZQ=="},
    }
    assert "uncertainties" in content[0]["text"]


@pytest.mark.asyncio
async def test_mimo_multimodal_photo_returns_verified_structure(monkeypatch):
    """MiMo 成功看图时，拍题链路返回题干和可审计的已知条件。

    此测试要防止的回归：接口返回 200 却把模型内容当普通文本处理，使图中
    条件、置信度和确认状态丢失。
    """
    class _Response:
        def raise_for_status(self):
            return None

        def json(self):
            return {
                "choices": [{"message": {"content": json.dumps({
                    "question_text": "已知函数 f(x)=x^2-1，求零点。",
                    "conditions": ["定义域为实数集"],
                    "diagram_entities": {},
                    "uncertainties": [],
                    "confidence": 0.96,
                })}}]
            }

    class _Http:
        async def post(self, url, *, headers, json):
            assert url.endswith("/chat/completions")
            assert json["model"] == fr.settings.mimo_vision_model  # 跟随配置（现为智谱 glm-4v-flash）
            return _Response()

    monkeypatch.setattr(fr.settings, "deepseek_api_key", "test-key")
    monkeypatch.setattr(fr, "_read_file_bytes", lambda _file: b"image-bytes")
    monkeypatch.setattr("app.providers.http.get_http", lambda: _Http())

    result = await fr._parse_image_mimo_vision(_file_obj())

    assert result is not None
    assert result["text"] == "已知函数 f(x)=x^2-1，求零点。"
    assert result["conditions"] == ["定义域为实数集"]
    assert result["needs_confirmation"] is False


@pytest.mark.asyncio
async def test_question_photo_prefers_mimo_before_ocr_fallback(monkeypatch):
    """拍题在 MiMo 识别成功时必须采用其结果，而不是先被 OCR 的错字污染。

    此测试要防止的回归：上传主链仍先执行 RapidOCR，即便 MiMo 已可可靠读图，
    导致公式和几何条件以低质量 OCR 文本进入课堂。
    """
    recognized = {
        "text": "已知 f(x)=x^2-1，求零点。",
        "conditions": ["x 属于 R"],
        "diagram_entities": {},
        "uncertainties": [],
        "confidence": 0.96,
        "needs_confirmation": False,
    }
    monkeypatch.setattr(fr, "_parse_image_mimo_vision", AsyncMock(return_value=recognized))
    fallback = AsyncMock(return_value="错误的 OCR 文本")
    monkeypatch.setattr(fr, "_dispatch_parse", fallback)

    content, confidence, parse_engine, evidence = await fr._parse_initial_file_content(
        _file_obj(), "rapidocr", "question_photo"
    )

    assert content == "已知 f(x)=x^2-1，求零点。"
    assert confidence == 0.96
    assert parse_engine == "mimo_vision"
    assert evidence["conditions"] == ["x 属于 R"]
    fallback.assert_not_awaited()


# ========== M2.2：/{file_id}/content 预签名端点（图片历史回显依赖） ==========

_content_engine = create_async_engine(_settings.database_url, poolclass=NullPool)
_content_session_factory = async_sessionmaker(
    _content_engine, class_=AsyncSession, expire_on_commit=False
)


async def _content_override_get_db():
    async with _content_session_factory() as session:
        try:
            yield session
            await session.commit()
        except Exception:
            await session.rollback()
            raise
        finally:
            await session.close()


_app.dependency_overrides[_get_db] = _content_override_get_db


class _FakePresignStorage:
    presign_expires = 900

    def __init__(self, url: str = "http://minio/bucket/k?sig=1", error: Exception | None = None):
        self._url = url
        self._error = error

    def presign_get(self, uri):
        if self._error:
            raise self._error
        return self._url


@pytest_asyncio.fixture
async def file_client():
    transport = ASGITransport(app=_app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        phone = f"137{str(_uuid.uuid4().int)[:8]}"
        await ac.post("/api/auth/sms-code", json={"phone": phone})
        resp = await ac.post("/api/auth/login", json={"phone": phone, "code": "123456"})
        data = resp.json()["data"]
        yield ac, data["token"], data["user"]["id"]


async def _seed_file(user_id: str) -> str:
    """落库一条文件记录（committed，端点可读），返回 file_id；用户不存在则补建行"""
    async with _content_session_factory() as s:
        from app.models.user import User

        owner = await s.get(User, _uuid.UUID(user_id))
        if owner is None:
            owner = User(id=_uuid.UUID(user_id), phone=f"136{str(_uuid.uuid4().int)[:8]}", nickname="")
            s.add(owner)
            await s.flush()
        f = File(
            user_id=_uuid.UUID(user_id),
            filename="t.png",
            file_type="image",
            mime="image/png",
            size_bytes=100,
            sha256="a" * 64,
            storage_uri=f"files/{user_id}/abc_t.png",
            status="parsed",
        )
        s.add(f)
        await s.commit()
        return str(f.id)


@pytest.mark.asyncio
async def test_content_endpoint_returns_presigned_url(file_client):
    """属主取 content → code=0 + 预签名 URL（历史图片回显链路）"""
    client, token, user_id = file_client
    file_id = await _seed_file(user_id)
    with patch.object(
        fr, "get_storage_for_user", new=AsyncMock(return_value=_FakePresignStorage())
    ):
        resp = await client.get(
            f"/api/files/{file_id}/content",
            headers={"Authorization": f"Bearer {token}"},
        )
    body = resp.json()
    assert body["code"] == 0
    assert body["data"]["url"] == "http://minio/bucket/k?sig=1"
    assert body["data"]["mime"] == "image/png"


@pytest.mark.asyncio
async def test_content_endpoint_storage_failure_structured_error(file_client):
    """存储未配置/不可用 → 结构化 50301（前端降级图标 chip），绝不裸 500"""
    client, token, user_id = file_client
    file_id = await _seed_file(user_id)
    boom = RuntimeError("对象存储未配置 endpoint")
    with patch.object(
        fr, "get_storage_for_user", new=AsyncMock(return_value=_FakePresignStorage(error=boom))
    ):
        resp = await client.get(
            f"/api/files/{file_id}/content",
            headers={"Authorization": f"Bearer {token}"},
        )
    assert resp.status_code == 200  # 结构化错误体，非 500
    assert resp.json()["code"] == 50301


@pytest.mark.asyncio
async def test_content_endpoint_not_owner_404(file_client):
    """越权取他人文件 → 40400（不泄露存在性）"""
    client, token, _ = file_client
    file_id = await _seed_file(str(_uuid.uuid4()))  # 别人的文件
    resp = await client.get(
        f"/api/files/{file_id}/content",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.json()["code"] == 40400


@pytest.mark.asyncio
async def test_development_photo_upload_falls_back_to_local_and_remains_attachable(file_client):
    client, token, _user_id = file_client
    auth = {"Authorization": f"Bearer {token}"}
    image_bytes = b"\x89PNG\r\n\x1a\nlocal-photo-fallback"

    with patch.object(fr.settings, "app_env", "development"), patch.object(
        fr.settings, "xingchen_enabled", False
    ), patch.object(fr, "_parse_image_rapidocr", new=AsyncMock(return_value=None)):
        initialized = await client.post(
            "/api/files/upload",
            json={
                "filename": "solution.png",
                "mime": "image/png",
                "size_bytes": len(image_bytes),
                "sha256": __import__("hashlib").sha256(image_bytes).hexdigest(),
                "multipart": False,
            },
            headers=auth,
        )
        assert initialized.json()["code"] == 0, initialized.text
        data = initialized.json()["data"]
        assert data["upload_url"].startswith(f"/api/files/{data['file_id']}/local-upload?")

        retried = await client.post(
            "/api/files/upload",
            json={
                "filename": "solution.png",
                "mime": "image/png",
                "size_bytes": len(image_bytes),
                "sha256": __import__("hashlib").sha256(image_bytes).hexdigest(),
                "multipart": False,
            },
            headers=auth,
        )
        assert retried.json()["data"]["deduplicated"] is False
        data = retried.json()["data"]

        uploaded = await client.put(
            data["upload_url"], content=image_bytes, headers={"Content-Type": "image/png"}
        )
        assert uploaded.status_code == 200, uploaded.text

        parsed = await client.post(
            f"/api/files/{data['file_id']}/parse",
            json={"purpose": "question_photo"},
            headers=auth,
        )
        assert parsed.json()["code"] == 0, parsed.text
        detail = await client.get(f"/api/files/{data['file_id']}", headers=auth)
        assert detail.json()["data"]["status"] == "parsed"
